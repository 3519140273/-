#!/usr/bin/env python3
"""生成重庆邮电大学实验报告册 .docx — 匹配参考文档格式"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

BASE = r'C:\Users\ASUS\Desktop\doubao报告'
BLACK = RGBColor(0, 0, 0)
doc = Document()

# ── 全局字体 ──
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(11)
style.font.color.rgb = BLACK

# ── 辅助函数 ──
def P(text=''):
    """普通段落，强制黑色"""
    p = doc.add_paragraph()
    if text:
        r = p.add_run(text)
        r.font.color.rgb = BLACK
        r.font.size = Pt(11)
    return p

def B(text, bold=False, size=11):
    """加粗段落"""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = BLACK
    return p

def I(text):
    """缩进段落"""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.color.rgb = BLACK
    r.font.size = Pt(11)
    p.paragraph_format.first_line_indent = Pt(22)
    return p

def code(text):
    """等宽代码行"""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    r.font.color.rgb = BLACK
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    return p

def pic(caption=''):
    """截图占位符"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'[pic]{(" " + caption) if caption else ""}')
    r.font.size = Pt(10)
    r.font.color.rgb = BLACK
    return p

def info_table(headers, rows):
    """插入表格"""
    t = doc.add_table(rows=1+len(rows), cols=len(headers), style='Table Grid')
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for r in c.paragraphs[0].runs:
            r.font.bold = True; r.font.size = Pt(10); r.font.color.rgb = BLACK
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(val)
            for r in c.paragraphs[0].runs:
                r.font.size = Pt(10); r.font.color.rgb = BLACK
    P()
    return t

def exp_header(exp_no, title, date, hours='4'):
    """每个实验的信息表"""
    info_table(
        ['姓名', '学号', '班级', '实验序号', '日期', '地点', '实验名称', '学时数', '指导教师', '成绩'],
        [['', '', '', f'实验{exp_no}', date, 'B608/609', title, hours, '', '']]
    )

# ============================================================
#                         封    面
# ============================================================
P()
pic('重庆邮电大学')
P()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('重 庆 邮 电 大 学')
r.font.size = Pt(22); r.font.bold = True; r.font.color.rgb = BLACK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('实 验 报 告 册')
r.font.size = Pt(20); r.font.bold = True; r.font.color.rgb = BLACK

P()
info_table(
    ['学年学期', '课程名称', '学号', '姓名', '电话'],
    [['2025-2026学年 第2学期', '机器人技术', '', '', '']]
)

P()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('重庆邮电大学教务处制')
r.font.size = Pt(10); r.font.color.rgb = BLACK

doc.add_page_break()

# ============================================================
#                      实验一
# ============================================================
exp_header('一', '常用开发和代码管理工具', '2026.5.14')
P()
B('（一）实验目的', bold=True)
I('掌握Roboware Studio、Git、Gazebo、Rviz、rqt和rosbag等ROS常用开发和代码管理工具的基本使用方法。了解如何在Roboware Studio中创建Catkin工作空间并编写ROS节点程序；如何使用Git进行版本控制并将代码托管到GitHub；如何使用Gazebo进行机器人仿真和Rviz进行可视化显示；如何使用rqt_graph查看ROS通信结构；如何使用rosbag录制和回放ROS消息数据。')

B('（二）实验要求', bold=True)
I('1. 在Roboware Studio中创建Catkin工作空间，编写ROS消息发布和接收节点，实现"Hello World"控制台输出。')
I('2. 创建GitHub仓库，配置SSH密钥，将上述程序文件推送到远程仓库。')
I('3. 运行Gazebo机器人仿真案例，使用Rviz显示传感器数据。')
I('4. 运行小乌龟例程，使用rqt_graph查看通信拓扑，使用rosbag录制并回放数据。')

B('实验设备', bold=True)
I('1. 硬件：计算机（Ubuntu 16.04/20.04操作系统）。')
I('2. 软件：ROS Noetic，Roboware Studio，Git，GitHub，Gazebo，Rviz，rqt，rosbag。')

B('实验原理', bold=True)
I('ROS（Robot Operating System）采用分布式节点通信架构，节点间通过话题（Topic）的发布/订阅模式进行通信。Catkin是ROS的编译构建系统，工作空间包含src（源码）、build（编译中间文件）、devel（环境脚本）三个目录。Git是分布式版本控制系统，通过SSH密钥连接GitHub远程仓库。Gazebo是三维物理仿真引擎，可加载机器人模型并模拟传感器和物理交互。Rviz是三维可视化工具，可订阅ROS话题并显示机器人模型、传感器数据等。rqt_graph以图形化方式展示节点和话题的通信关系。rosbag可录制ROS消息到bag文件，支持离线回放，便于调试和数据分析。')

B('实验步骤/方法', bold=True)

B('（1）Roboware Studio编写ROS Hello World程序')
I('① 创建Catkin工作空间：')
code('mkdir -p ~/catkin_ws/src')
code('cd ~/catkin_ws')
code('catkin_make')
code('source devel/setup.bash')
pic()
I('② 创建Package。在src目录下创建hello_ros包，依赖roscpp和std_msgs：')
I('③ 编写源代码。在Roboware Studio中创建两个Python节点文件：')
I('talker.py（消息发布者）：')
code('#!/usr/bin/env python3')
code('import rospy')
code('from std_msgs.msg import String')
code('')
code('rospy.init_node("talker_node")')
code('pub = rospy.Publisher("chatter", String, queue_size=10)')
code('rate = rospy.Rate(10)')
code('while not rospy.is_shutdown():')
code('    msg = "Hello World !"')
code('    rospy.loginfo(msg)')
code('    pub.publish(msg)')
code('    rate.sleep()')
I('listener.py（消息订阅者）：')
code('#!/usr/bin/env python3')
code('import rospy')
code('from std_msgs.msg import String')
code('')
code('def callback(data):')
code('    rospy.loginfo("I heard: %s", data.data)')
code('')
code('rospy.init_node("listener_node")')
code('rospy.Subscriber("chatter", String, callback)')
code('rospy.spin()')
I('④ 编译运行。执行catkin_make编译，然后依次运行：')
code('roscore                      # 终端1：启动ROS Master')
code('rosrun hello_ros talker.py   # 终端2：启动发布者')
code('rosrun hello_ros listener.py # 终端3：启动订阅者')
I('终端输出"Hello World !"，验证消息收发成功。')
pic('Hello World运行截图')

B('（2）Git配置SSH并上传至GitHub')
I('注册GitHub账号，使用ssh-keygen生成SSH密钥，将公钥添加到GitHub设置中。将hello_ros项目初始化为Git仓库并推送：')
code('cd ~/catkin_ws/src/hello_ros')
code('git init')
code('git add .')
code('git commit -m "first commit"')
code('git remote add origin git@github.com:用户名/hello_ros.git')
code('git push -u origin main')
I('验证：登录GitHub查看仓库，确认代码文件已成功上传。')
pic('GitHub仓库截图')

B('（3）Gazebo仿真与Rviz可视化')
I('运行实验室提供的机器人仿真案例：')
code('roslaunch robot_sim_demo robot_spawn.launch')
I('Gazebo启动后显示XBot机器人仿真场景，Rviz同步显示机器人视角下的传感器数据（激光雷达点云、摄像头图像等），验证仿真环境正常工作。')
pic('Gazebo+Rviz仿真截图')

B('（4）rqt_graph查看通信结构与rosbag数据录制回放')
I('启动小乌龟例程：')
code('roscore')
code('rosrun turtlesim turtlesim_node')
code('rosrun turtlesim turtle_teleop_key')
I('打开rqt_graph查看当前ROS系统中节点和话题的通信关系：')
code('rqt_graph')
I('使用键盘控制小乌龟运动，同时录制话题数据：')
code('rosbag record -a')
I('停止录制后，关闭turtlesim节点，使用rosbag回放数据：')
code('rosbag play <录制的bag文件名>')
I('rqt_graph中可清晰看到/turtlesim节点订阅/cmd_vel话题、发布/pose话题的通信结构；rosbag成功录制并回放了所有话题数据。')
pic('rqt_graph截图')

B('结果与分析', bold=True)
I('（1）遇到的问题')
I('Catkin工作空间首次编译时出现权限错误，通过chmod命令修改文件权限后解决。Git推送时SSH连接被拒绝，检查发现公钥未正确添加到GitHub，重新添加后正常推送。rosbag录制时忘记先source环境变量，导致命令未找到，刷新环境后正常。')
I('（2）心得')
I('通过本实验熟悉了ROS开发的基本工具链：Roboware Studio提供IDE开发环境，Git管理代码版本，Gazebo和Rviz实现仿真与可视化，rqt_graph和rosbag辅助调试分析。这些工具相互配合，构成了完整的ROS开发流程，为后续实验奠定了基础。各工具的熟练使用对ROS开发效率有重要影响，特别是Git版本管理可有效避免代码丢失和混乱。')

doc.add_page_break()

# ============================================================
#                      实验二
# ============================================================
exp_header('二', '机器人URDF模型建立、控制及Gazebo仿真', '2026.5.29')
P()
B('（一）实验目的', bold=True)
I('掌握URDF（Unified Robot Description Format）建立机器人三维模型的方法，学习使用ArbotiX和Rviz进行模型显示与运动控制，了解ros_control机器人控制框架，实现在Gazebo仿真环境中控制机器人运动并获取传感器数据。')

B('（二）实验要求', bold=True)
I('1. 创建四轮小车URDF模型：后轮驱动、前轮转向，车顶安装摄像头。使用XML描述机器人的link（刚体）和joint（关节）属性。')
I('2. 基于ArbotiX和Rviz控制小车运动：编写launch文件加载模型，使用joint_state_publisher_gui调节关节角度，在Rviz中显示机器人模型。')
I('3. 配置ros_control控制器，在Gazebo中仿真小车运动，订阅并显示摄像头图像话题。')

B('实验设备', bold=True)
I('1. 硬件：计算机（Ubuntu 20.04操作系统）。')
I('2. 软件：ROS Noetic，Rviz，Gazebo，ArbotiX，ros_control，xacro，robot_state_publisher，rqt_image_view。')

B('实验原理', bold=True)
I('URDF是ROS中描述机器人模型的标准XML格式。模型由<link>和<joint>两类元素组成：<link>定义刚体部件，包含visual（视觉）、collision（碰撞）、inertial（惯性）属性；<joint>定义部件间的连接关系，支持fixed（固定）、continuous（连续旋转）、revolute（限位旋转）等类型。')
I('ArbotiX是ROS的差速驱动控制器，订阅/cmd_vel话题（Twist消息）并计算左右轮速度，发布/joint_states话题。Rviz通过robot_state_publisher订阅/joint_states并发布TF坐标变换，从而可视化机器人模型各部件的位置和姿态。')
I('ros_control是ROS的机器人控制框架，包含控制器管理器（Controller Manager）、PID控制器、关节状态接口（JointStateInterface）等模块，通过YAML配置文件定义各关节的控制器类型和PID参数。Gazebo集成URDF模型后，通过libgazebo_ros_diff_drive等插件实现物理仿真，通过libgazebo_ros_camera插件发布模拟摄像头图像。')

B('实验步骤/方法', bold=True)

B('（1）创建四轮小车URDF模型')
I('在ROS工作空间src目录下创建car_description包，在urdf目录中编写car.urdf文件，定义车身（base_link）、左后轮（left_rear_wheel）、右后轮（right_rear_wheel）、左前轮（left_front_wheel）、右前轮（right_front_wheel）和摄像头（camera_link）六个link及其关节。关键模型参数如下：')
I('车身尺寸：0.5m×0.3m×0.1m（蓝色箱体），质量10kg。')
I('车轮：半径0.08m、宽度0.04m的圆柱体，后轮间距0.34m，轴距0.30m。后轮joint类型continuous（驱动轮），前轮joint类型continuous（转向轮）。')
I('摄像头：0.05m×0.05m×0.05m的绿色箱体，通过fixed joint固定在车身顶部（z=0.15m）。')
I('Gazebo插件配置：差速驱动插件（libgazebo_ros_diff_drive.so）控制左右后轮，wheelSeparation=0.34，wheelDiameter=0.16；摄像头插件（libgazebo_ros_camera.so）发布/camera/image_raw话题，分辨率640×480，更新率30Hz。')
pic('URDF模型代码结构截图')

B('（2）ArbotiX + Rviz显示与控制')
I('编写launch文件car_display.launch：')
code('<launch>')
code('  <param name="robot_description" command="$(find xacro)/xacro $(find car_description)/urdf/car.urdf"/>')
code('  <node pkg="robot_state_publisher" type="robot_state_publisher" name="robot_state_publisher"/>')
code('  <node pkg="joint_state_publisher_gui" type="joint_state_publisher_gui" name="joint_state_publisher_gui"/>')
code('  <node pkg="rviz" type="rviz" name="rviz"/>')
code('</launch>')
I('运行roslaunch car_description car_display.launch，在Rviz中添加RobotModel显示项，调节joint_state_publisher_gui滑块控制各关节转动，验证模型结构正确。')
pic('Rviz中URDF模型显示截图')

B('（3）ros_control配置与Gazebo仿真')
I('创建config/controllers.yaml配置文件，定义后轮速度控制器和前轮位置控制器：')
code('left_rear_wheel_joint:')
code('  type: effort_controllers/JointVelocityController')
code('  joint: left_rear_wheel_joint')
code('  pid: {p: 100.0, i: 0.0, d: 0.0}')
code('right_rear_wheel_joint:')
code('  type: effort_controllers/JointVelocityController')
code('  joint: right_rear_wheel_joint')
code('  pid: {p: 100.0, i: 0.0, d: 0.0}')
I('编写car_gazebo.launch启动文件，依次启动Gazebo空世界、加载URDF模型、spawn小车、加载控制器参数、启动Rviz。运行后Gazebo显示小车模型，通过键盘teleop节点发送/cmd_vel指令控制小车运动，使用rqt_image_view订阅/camera/image_raw话题查看模拟摄像头图像。')
pic('Gazebo仿真截图')
pic('摄像头图像截图')

B('结果与分析', bold=True)
I('（1）遇到的问题')
I('首次加载URDF模型时Rviz中模型不显示，检查发现robot_description参数未正确加载xacro文件，在launch文件中添加$(find xacro)/xacro前缀后解决。Gazebo中前轮转向功能不工作，排查发现controllers.yaml中前轮joint名称与URDF中定义不一致，统一命名后正常。PID参数初始设置偏小导致车轮响应迟缓，增大P值至100后响应速度改善。')
I('（2）心得')
I('通过本实验深入理解了URDF建模方法和link/joint的设计原则。对比了ArbotiX（纯运动学仿真）与Gazebo（物理仿真）两种仿真方式：ArbotiX配置简单，适合快速验证模型结构；Gazebo需要配置ros_control和物理参数，但仿真效果更接近真实机器人。URDF模型的完整性和准确性直接影响后续导航和控制效果，建立模型时需仔细核对坐标系和各部件的相对位置。')

doc.add_page_break()

# ============================================================
#                      实验三
# ============================================================
exp_header('三', 'ROS语音对话实验', '2026.6.6')
P()
B('（一）实验目的', bold=True)
I('学习安装和使用PocketSphinx语音识别功能包，掌握离线语音识别的基本方法；学习使用科大讯飞SDK实现在线语音听写，了解云端语音识别API的调用方式；实现通过语音命令控制ROS机器人（小海龟）的运动，理解语音识别→语义解析→运动控制的完整流程。')

B('（二）实验要求', bold=True)
I('1. 安装PocketSphinx功能包及其中文声学模型，编写ROS语音识别节点，实现麦克风语音到文本的离线转换。')
I('2. 集成科大讯飞语音听写SDK（iat），使用WebSocket和HMAC-SHA256鉴权，实现高精度在线语音识别。')
I('3. 编写语音控制节点，将识别文本解析为运动指令，通过/cmd_vel话题控制turtlesim小海龟的前进、后退、左转、右转和停止。')

B('实验设备', bold=True)
I('1. 硬件：计算机（Ubuntu 20.04操作系统），麦克风。')
I('2. 软件：ROS Noetic，PocketSphinx，科大讯飞SDK（iat语音听写API），pyaudio，websocket-client，speech_recognition库，turtlesim。')

B('实验原理', bold=True)
I('PocketSphinx是CMU开发的开源语音识别引擎，基于HMM（隐马尔可夫模型）声学模型，支持离线中英文识别，不需要网络连接。中文识别需要额外安装zh-cn声学模型、语言模型和发音字典。识别流程为：麦克风采集音频→特征提取（MFCC）→声学模型匹配→语言模型解码→输出文本。')
I('科大讯飞语音听写API采用WebSocket长连接方案：客户端通过HMAC-SHA256算法生成鉴权签名，建立wss连接后先发送参数帧（包含app_id、语言、采样率等），然后持续发送PCM音频数据帧，服务器实时返回中间识别结果和最终结果。相比PocketSphinx，云端识别准确率更高、支持连续语音和更长句子。')
I('语音控制流程：识别文本→关键词匹配（如"前进"→forward）→生成Twist消息（设置linear.x或angular.z）→发布到/cmd_vel话题→turtlesim_node订阅并执行运动。')

B('实验步骤/方法', bold=True)

B('（1）PocketSphinx离线语音识别')
I('安装PocketSphinx及中文模型（cmusphinx-zh-cn-5.2），使用Python的speech_recognition库封装识别功能：')
code('import rospy, speech_recognition as sr')
code('from std_msgs.msg import String')
code('')
code('rospy.init_node("sphinx_demo")')
code('pub = rospy.Publisher("/voice/recognized", String, queue_size=10)')
code('rec = sr.Recognizer()')
code('mic = sr.Microphone()')
code('')
code('with mic as source:')
code('    rec.adjust_for_ambient_noise(source, duration=1)  # 降噪')
code('    audio = rec.listen(source, timeout=5)')
code('text = rec.recognize_sphinx(audio)  # 离线语音识别')
code('pub.publish(String(data=text))')
I('运行节点后，对着麦克风说出简单的英文或中文词汇，终端打印识别结果并发布到/voice/recognized话题。PocketSphinx对短词汇识别效果较好，支持数字、简单指令等。')
pic('PocketSphinx识别结果截图')

B('（2）科大讯飞SDK在线语音识别与控制')
I('注册科大讯飞开放平台账号，创建语音听写应用获取APP_ID、API_KEY、API_SECRET。通过WebSocket连接讯飞云端API：')
code('# HMAC-SHA256 鉴权')
code('sig_raw = f"host: {HOST}\\ndate: {date}\\nGET /v2/iat HTTP/1.1"')
code('sig = base64.b64encode(hmac.new(API_SECRET.encode(), sig_raw.encode(), hashlib.sha256).digest()).decode()')
code('ws_url = f"wss://{HOST}/v2/iat?authorization={quote(auth)}&date={date}&host={HOST}"')
code('')
code('# 参数帧（需先发送，再发音频数据）')
code('params = {"common":{"app_id":APP_ID}, "business":{"language":"zh_cn","domain":"iat","accent":"mandarin"}, "data":{"status":0,"format":"audio/L16;rate=16000","encoding":"raw","audio":""}}')
code('')
code('# 音频帧（PCM 16kHz 16bit 单声道，base64编码）')
code('chunk = stream.read(3200)  # pyaudio读取麦克风')
code('frame = {"data":{"status":1,"format":"audio/L16;rate=16000","encoding":"raw","audio":base64.b64encode(chunk).decode()}}')
I('语音命令解析与执行：')
code('def parse(text):')
code('    if "前进" in text or "向前" in text: return "forward"')
code('    if "后退" in text or "向后" in text: return "backward"')
code('    if "左转" in text: return "left"')
code('    if "右转" in text: return "right"')
code('    if "停止" in text: return "stop"')
code('')
code('def execute(cmd):')
code('    t = Twist()')
code('    if cmd == "forward":   t.linear.x = 1.0')
code('    elif cmd == "backward": t.linear.x = -1.0')
code('    elif cmd == "left":     t.angular.z = 1.0')
code('    elif cmd == "right":    t.angular.z = -1.0')
code('    cmd_pub.publish(t)')
code('    rospy.sleep(1.0)')
code('    cmd_pub.publish(Twist())  # 停止')
I('启动voice_control.launch（同时启动turtlesim_node和voice_control节点），对着麦克风说出"前进""后退""左转""右转""停止"等命令词，小海龟按语音指令运动。')
pic('语音控制小海龟截图')

B('结果与分析', bold=True)
I('（1）遇到的问题')
I('PocketSphinx中文模型安装后识别率很低，排查发现声学模型路径配置错误，修正hmm/lm/dic路径后识别率提升。科大讯飞SDK连接时出现401鉴权错误，检查发现API_SECRET末尾有多余空格，修剪后正常。WebSocket连接不稳定时出现"Connection reset"异常，添加重试机制和超时处理后改善。麦克风增益设置过高导致削波失真，通过adjust_for_ambient_noise自动调节后解决。')
I('（2）心得')
I('通过本实验理解了语音识别的基本原理和技术路线。PocketSphinx适合网络受限的离线场景和简单词汇识别；科大讯飞SDK调用云端模型，识别准确率显著优于离线方案，适合复杂语音交互场景。ROS的话题通信机制使语音识别节点与控制节点解耦，便于替换不同识别方案。从语音采集到最终控制执行的完整流程中，各环节（采集质量、识别准确率、命令解析鲁棒性）都会影响最终效果，需要综合优化。')

doc.add_page_break()

# ============================================================
#                      实验四
# ============================================================
exp_header('四', '人脸检测和自主导航实验', '2026.6.12')
P()
B('（一）实验目的', bold=True)
I('学习安装和使用OpenCV库，掌握基于Haar级联分类器的人脸检测方法；学习ROS SLAM建图技术（gmapping），掌握地图构建与保存的方法；学习ROS机器人自主导航框架（move_base），理解全局规划、局部规划和AMCL定位的协作关系，实现在建立的地图上进行自主导航。')

B('（二）实验要求', bold=True)
I('1. 安装OpenCV库，编写ROS人脸检测节点，订阅摄像头图像话题，使用Haar Cascade分类器检测人脸并用矩形框标注。')
I('2. 使用gmapping算法进行激光SLAM建图，遥控机器人在Gazebo仿真环境中运动，构建环境栅格地图并保存。')
I('3. 配置AMCL定位和move_base导航框架（全局代价地图、局部代价地图、DWA局部规划器），实现基于地图的自主导航功能。')

B('实验设备', bold=True)
I('1. 硬件：计算机（Ubuntu 20.04操作系统）。')
I('2. 软件：ROS Noetic，OpenCV，cv_bridge，gmapping，amcl，move_base，Navfn全局规划器，DWA局部规划器，Gazebo。')

B('实验原理', bold=True)
I('OpenCV人脸检测基于Haar级联分类器（Haar Cascade Classifier）。Haar特征是一种矩形特征模板，通过积分图快速计算；AdaBoost算法训练强分类器；级联结构将多个分类器串联，逐级筛选，在前几级快速排除非人脸区域以提高检测速度。检测流程：图像灰度化→直方图均衡化→多尺度滑动窗口扫描→Haar特征提取→级联分类器判定。')
I('gmapping（Grid-based SLAM）基于Rao-Blackwellized粒子滤波算法。每个粒子携带一份地图估计和机器人位姿，通过激光雷达数据更新粒子权重，重采样维持粒子多样性。关键参数包括：粒子数（particles）决定估计精度和计算量，地图更新间隔（map_update_interval）控制建图频率，最大激光距离（maxRange）限制数据使用范围。')
I('自主导航采用move_base框架：全局规划器（Navfn）基于静态地图使用Dijkstra/A*算法计算全局路径；局部规划器（DWA，Dynamic Window Approach）在速度空间中采样，评估每个速度组合的轨迹得分（朝向目标、避开障碍、速度最大化），选择最优轨迹执行；AMCL（自适应蒙特卡洛定位）使用粒子滤波估计机器人在已知地图中的位姿。代价地图分为全局（global_costmap，基于静态地图，全局坐标系）和局部（local_costmap，滚动窗口，里程计坐标系）两层，均包含静态层、障碍物层和膨胀层。')

B('实验步骤/方法', bold=True)

B('（1）OpenCV人脸检测')
I('创建face_detector节点，订阅Gazebo仿真摄像头图像话题/camera/rgb/image_raw，使用CvBridge将ROS Image消息转换为OpenCV格式，Haar Cascade检测人脸并标注：')
code('import rospy, cv2')
code('from sensor_msgs.msg import Image')
code('from cv_bridge import CvBridge')
code('')
code('class FaceDetector:')
code('    def __init__(self):')
code('        self.bridge = CvBridge()')
code('        cascade = cv2.data.haarcascades + "haarcascade_frontalface_default.xml"')
code('        self.face_cascade = cv2.CascadeClassifier(cascade)')
code('        rospy.Subscriber("/camera/rgb/image_raw", Image, self.callback)')
code('    ')
code('    def callback(self, msg):')
code('        cv_image = self.bridge.imgmsg_to_cv2(msg, "bgr8")')
code('        gray = cv2.cvtColor(cv_image, cv2.COLOR_BGR2GRAY)')
code('        gray = cv2.equalizeHist(gray)  # 直方图均衡化增强对比度')
code('        faces = self.face_cascade.detectMultiScale(')
code('            gray, scaleFactor=1.1, minNeighbors=5, minSize=(30,30))')
code('        for (x, y, w, h) in faces:')
code('            cv2.rectangle(cv_image, (x,y), (x+w,y+h), (0,255,0), 2)')
code('        cv2.imshow("Face Detection", cv_image)')
I('运行roslaunch my_robot experiment1_face_detect.launch，Gazebo仿真环境中的模拟人脸图像被检测并标注。')
pic('人脸检测结果截图')

B('（2）SLAM建图')
I('编写gmapping.launch启动gmapping节点，配置核心参数：base_frame=base_footprint，odom_frame=odom，particles=80，map_update_interval=1.0s，maxRange=15m，occ_thresh=0.25。')
I('启动Gazebo仿真环境和gmapping建图节点：')
code('roslaunch my_robot gazebo.launch')
code('roslaunch my_robot gmapping.launch')
I('使用键盘teleop遥控机器人在仿真房间中运动（四个墙壁、多个障碍物），gmapping实时更新栅格地图。通过Rviz可视化建图过程和激光雷达扫描数据。建图完成后保存地图：')
code('rosrun map_server map_saver -f ~/catkin_ws/src/my_robot/maps/my_map')
I('生成my_map.pgm（栅格地图图像）和my_map.yaml（地图元数据：分辨率0.05m/pixel，原点坐标，占用/空闲阈值）。')
pic('SLAM建图结果截图')

B('（3）自主导航')
I('配置AMCL定位节点（amcl_navigation.launch）：odom_model_type=diff（差速模型），min_particles=500，max_particles=5000，laser_model_type=likelihood_field。')
I('配置move_base导航参数：')
code('# costmap_common.yaml: 障碍物检测范围3m，robot_radius=0.26m，膨胀半径0.55m')
code('# global_costmap.yaml: 全局坐标系map，使用静态地图层，不滚动')
code('# local_costmap.yaml: 局部坐标系odom，滚动窗口4m×4m，分辨率0.05m')
code('# move_base.yaml: 全局规划器navfn/NavfnROS，局部规划器dwa_local_planner/DWAPlannerROS')
code('# dwa_planner.yaml: max_vel_x=0.8m/s，max_vel_theta=2.0rad/s，')
code('#   xy_goal_tolerance=0.15m，yaw_goal_tolerance=0.15rad')
I('启动导航：')
code('roslaunch my_robot amcl_navigation.launch')
I('在Rviz中使用"2D Nav Goal"工具点击目标位置和朝向，机器人自动规划全局路径（绿色线）并以DWA计算的局部轨迹（黄色线）执行，实时避开障碍物到达目标点。')
pic('自主导航截图')

B('结果与分析', bold=True)
I('（1）遇到的问题')
I('人脸检测节点启动后无法加载Haar Cascade文件，原因是cv2.data.haarcascades路径在高版本OpenCV中位置变化，使用绝对路径替代后解决。gmapping建图时粒子数设置过低（20）导致建图质量差、出现重影，增加到80后改善明显。导航时机器人在窄通道中出现振荡，调整inflation_radius从0.3m增至0.55m、减小xy_goal_tolerance至0.15m后效果改善。AMCL初始位姿偏差较大时收敛缓慢，通过设置初始位姿估计加速收敛。')
I('（2）心得')
I('通过本实验掌握了ROS视觉处理和自主导航的完整技术栈。OpenCV Haar级联分类器在正面人脸场景下检测效果良好，但侧脸和遮挡情况下漏检较多；深度学习方法（如MTCNN、RetinaFace）可进一步提升检测鲁棒性。gmapping适合中小规模室内环境建图，粒子数和激光数据质量直接影响建图精度。move_base框架将全局规划、局部规划和定位模块化整合，各模块参数独立可调，便于针对不同场景优化。该技术栈可直接迁移至真实机器人平台，实现未知环境的自主探索与导航。')

doc.add_page_break()

# ============================================================
#                         保    存
# ============================================================
output_path = os.path.join(BASE, '实验报告册_新版.docx')
doc.save(output_path)
print(f'报告已生成: {output_path}')
