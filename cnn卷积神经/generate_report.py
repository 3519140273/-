"""
生成CNN课程报告Word文档
黑色字体，中文宋体+英文TNR 小四，客观学术风格
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


def _set_style_font(style, cn, en, size, bold=False, color=None):
    """直接修改Word样式（确保字体生效）"""
    font = style.font
    font.name = en
    font.size = size
    font.bold = bold
    if color:
        font.color.rgb = color
    # 通过XML设置中文字体
    rPr = style.element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), en)
    rFonts.set(qn('w:eastAsia'), cn)
    rFonts.set(qn('w:hAnsi'), en)
    rPr.insert(0, rFonts)


def set_run_font(run, cn='宋体', en='Times New Roman', size=Pt(12), bold=False):
    run.font.size = size
    run.bold = bold
    run.font.name = en
    run.font.color.rgb = None
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), en)
    rFonts.set(qn('w:eastAsia'), cn)
    rFonts.set(qn('w:hAnsi'), en)
    rPr.insert(0, rFonts)


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    sizes = {1: Pt(16), 2: Pt(14), 3: Pt(13)}
    for run in h.runs:
        set_run_font(run, '黑体', 'Times New Roman', sizes.get(level, Pt(13)), bold=True)


def body(doc, text, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    # 段落级别字体设置
    pf = p.paragraph_format
    r = p.add_run(text)
    set_run_font(r)
    return p


def fig_placeholder(doc, caption, filename):
    """图片占位"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run(f'[插入图片: {filename}]')
    set_run_font(r, '宋体', 'Times New Roman', Pt(11), bold=True)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(
        '┌──────────────────────────────────────────────────┐\n'
        '│                                                  │\n'
        f'│   {filename:<46s}│\n'
        '│                                                  │\n'
        '└──────────────────────────────────────────────────┘'
    )
    set_run_font(r2, '宋体', 'Courier New', Pt(8))

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(caption)
    set_run_font(r3, '宋体', 'Times New Roman', Pt(10))
    p3.paragraph_format.space_after = Pt(8)


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Table Grid')
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); set_run_font(r, '宋体', 'Times New Roman', Pt(10), bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val)); set_run_font(r, '宋体', 'Times New Roman', Pt(10))
    doc.add_paragraph()


def load_history(path='results/training_history.txt'):
    if not os.path.exists(path):
        return []
    with open(path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    data = []
    for line in lines[2:]:
        parts = [p.strip() for p in line.strip().split('|')]
        if len(parts) >= 5:
            data.append(parts)
    return data


def generate():
    doc = Document()

    # ===== 全局修改文档样式 =====
    # Normal样式 -> 正文: 宋体/TNR 小四
    _set_style_font(doc.styles['Normal'], '宋体', 'Times New Roman', Pt(12))
    # Heading样式 -> 标题: 黑体/TNR
    for lvl in [1, 2, 3]:
        size_map = {1: Pt(16), 2: Pt(14), 3: Pt(13)}
        _set_style_font(doc.styles[f'Heading {lvl}'], '黑体', 'Times New Roman', size_map[lvl], bold=True)

    for s in doc.sections:
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(2.8)
        s.right_margin = Cm(2.5)

    # ======== 封面 ========
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('基于卷积神经网络的图形分类实验报告')
    set_run_font(r, '黑体', 'Times New Roman', Pt(22), bold=True)
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('—— 圆形与方形二分类任务')
    set_run_font(r, '黑体', 'Times New Roman', Pt(16), bold=True)
    for _ in range(6):
        doc.add_paragraph()
    for line in ['院系：__________', '专业：__________', '姓名：__________', '学号：__________', '日期：2026年6月']:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line); set_run_font(r, '宋体', 'Times New Roman', Pt(14))
    doc.add_page_break()

    # ======== 一、引言 ========
    heading(doc, '一、引言', 1)

    body(doc,
        '卷积神经网络（CNN）自LeNet-5[1]提出以来，已成为图像分类任务的基础模型。'
        '其核心思想是通过卷积核在图像上滑动提取局部特征，逐层组合为高层语义表示，'
        '最后经全连接层输出分类概率。本实验不采用常见的MNIST或CIFAR-10等公开数据集，'
        '而是用程序直接生成几何形状图片：圆形和方形。选择自行生成数据的原因有三：一是'
        '避免因网络条件限制无法下载大型数据集；二是可以精确控制数据的各项参数（噪声、'
        '位移、旋转等），便于进行消融分析；三是任务难度可控，能够保证模型收敛，从而将'
        '实验重心放在理解CNN的工作原理上。')

    body(doc,
        '实验的核心目标是验证CNN对几何形状轮廓特征的提取能力。圆形和方形的区分依据主要'
        '在于边缘的曲率特征——圆的曲率恒定，方形的曲率集中在四个角点。这个任务在理论上'
        '对CNN是友好的，因为卷积核天然适合检测边缘和角点。本报告将记录实现过程、模型'
        '结构选择理由、实验结果及存在的不足。')

    # ======== 二、实现方案 ========
    heading(doc, '二、实现方案', 1)

    heading(doc, '2.1 整体技术路线', 2)
    body(doc,
        '本实验基于PyTorch框架实现。选择PyTorch而非TensorFlow的原因是PyTorch的动态计算图'
        '机制更便于调试，且社区文档对新用户更友好。实验流程为：数据生成 → 模型构建 → '
        '训练与验证 → 评估与可视化。所有代码以模块化方式组织：model.py负责模型定义，'
        'utils.py负责数据生成与可视化，train.py负责训练流程。')

    heading(doc, '2.2 数据生成方案', 2)
    body(doc,
        '数据集完全由numpy和Python标准库生成，无需网络连接。核心参数如下表：')

    table(doc,
        ['参数', '取值', '说明'],
        [
            ['图像尺寸', '28×28', '与经典MNIST保持一致，便于参照'],
            ['通道数', '1（灰度）', '简化输入，降低计算量'],
            ['训练样本数', '5000', '圆形2500 + 方形2500，保证类别均衡'],
            ['验证样本数', '1000', '不参与训练，用于评估泛化能力'],
            ['噪声类型', '高斯噪声 σ=0.25', '模拟实际拍摄中的随机干扰'],
            ['形状位置', '中心随机偏移', 'x, y ∈ [7, 21]，避免形状贴边'],
            ['形状尺寸', '随机缩放', '半径/边长 ∈ [3.4, 9.0] 像素'],
            ['方形旋转', '随机 0°~45°', '增加方形样本多样性'],
        ]
    )

    body(doc,
        '生成圆形的算法：对于每个像素点计算其到圆心的欧氏距离d，若|d - r| < thickness则'
        '赋值为亮度值，边缘处做线性衰减以避免锯齿。生成方形的算法：对像素坐标做随机角度'
        '的旋转变换后，在旋转坐标系下判断是否位于矩形区域内，边缘同样做渐变处理。两种形状'
        '的亮度值在0.6~1.0之间随机采样，防止模型仅靠亮度做判断。最后叠加均值为0、标准差'
        '为0.25的高斯噪声。')

    fig_placeholder(doc, '图1 生成数据样本（上排: 圆形, 下排: 方形）', 'samples.png')

    body(doc,
        '图1展示了16个随机生成的训练样本。可以看到，噪声水平为0.25时，部分样本的轮廓'
        '已经比较模糊，这对CNN的特征提取能力构成一定挑战。但圆形与方形的整体形态差异'
        '仍然可辨。')

    heading(doc, '2.3 模型结构设计', 2)
    body(doc,
        '模型结构的设计遵循从简单到有效的原则。最初尝试了2层卷积的结构：Conv(1→8) → '
        'Pool → Conv(8→16) → Pool → FC(16×7×7 → 2)。该模型参数量仅约6,000个，在'
        '噪声为0时的准确率可达99%以上，但当噪声增加到0.25时准确率降至约85%，表现不稳定。')

    body(doc,
        '分析原因：卷积核数量过少导致特征提取不充分，在高噪声条件下无法可靠地区分形状。'
        '因此将通道数提升至{16, 32, 64}，并引入Batch Normalization（BN）层和Dropout。'
        'BN层的作用是稳定每层输入的分布，缓解内部协变量偏移（Internal Covariate Shift），'
        '加速收敛。Dropout在训练时随机屏蔽40%的神经元，等效于集成多个子网络，用于抑制'
        '过拟合。最终模型结构如下：')

    table(doc,
        ['层', '操作', '输出尺寸', '参数量', '备注'],
        [
            ['Conv1', '1→16, k=3, p=1', '16×28×28', '160', ''],
            ['BN1 + ReLU + Pool', 'MaxPool 2×2', '16×14×14', '32', 'BN: 16×2个参数'],
            ['Conv2', '16→32, k=3, p=1', '32×14×14', '4,640', ''],
            ['BN2 + ReLU + Pool', 'MaxPool 2×2', '32×7×7', '64', ''],
            ['Conv3', '32→64, k=3, p=0', '64×5×5', '18,496', 'k=3, p=0时7→5'],
            ['BN3 + ReLU + Pool', 'MaxPool 2×2', '64×2×2', '128', ''],
            ['Flatten', '向量化', '256', '0', '64×2×2=256'],
            ['FC1 + ReLU + Dropout', '256→128, p=0.4', '128', '32,896', '含bias 128'],
            ['FC2', '128→2', '2', '258', '含bias 2'],
        ]
    )

    body(doc, '模型总参数量为56,674个，全部参与训练（无冻结层）。')

    body(doc,
        '关于第三个卷积层使用padding=0的设计：输入尺寸为7×7，卷积核3×3，padding=0时输出'
        '为5×5，再经2×2池化得到2×2。如果使用padding=1维持7×7，池化后将得到3×3（向下取整），'
        '后续全连接层输入变为64×3×3=576。两种方案的参数量和效果差异不大，选择padding=0的'
        '方案仅是为了验证不padding时CNN仍能正常工作。')

    # ======== 三、实验设置与结果 ========
    heading(doc, '三、实验设置与结果', 1)

    heading(doc, '3.1 训练配置', 2)
    table(doc,
        ['配置项', '取值', '选择理由'],
        [
            ['优化器', 'Adam, lr=0.001', '自适应学习率，收敛快，超参数不敏感'],
            ['学习率调度', 'StepLR, step=5, γ=0.5', '每5步减半，后期精细化搜索'],
            ['Batch Size', '64', '折中选择，兼顾训练速度和梯度稳定性'],
            ['Epochs', '15', '15轮后Loss和Acc均趋于平稳'],
            ['损失函数', 'CrossEntropyLoss', '二分类标准交叉熵，内置Softmax'],
            ['评估指标', 'Accuracy/Precision/Recall/F1', '适用于类别均衡的二分类任务'],
        ]
    )

    heading(doc, '3.2 训练过程', 2)

    hist = load_history()
    # 选取几个代表性epoch
    key_epochs = [0, 2, 4, 7, 10, 14]
    key_data = [hist[i] for i in key_epochs if i < len(hist)]
    if key_data:
        table(doc,
            ['Epoch', 'Train Loss', 'Val Loss', 'Train Acc(%)', 'Val Acc(%)'],
            key_data
        )

    body(doc,
        '上表记录了训练过程中Loss和准确率的变化。可以看到，模型在第1个epoch时已接近收敛'
        '（训练准确率99.16%，验证准确率99.40%），说明圆形与方形的区分特征比较明显，CNN'
        '能快速学习到。后续epoch中Loss继续缓慢下降但波动增大，验证Loss在0.02~0.04之间'
        '振荡，最终在第11个epoch达到最高验证准确率99.40%。由于噪声水平较高，后期训练'
        '出现了更多的Loss尖峰，表明模型在部分噪声样本上的分类边界不够稳定。')

    fig_placeholder(doc, '图3 训练与验证的Loss曲线和Accuracy曲线', 'training_curves.png')

    body(doc,
        '从图3的Loss曲线观察，训练Loss始终低于验证Loss，这是正常现象——验证Loss计算时'
        'Dropout关闭且无数据增强带来的随机性。两者差距不大（约0.002~0.005），未见明显'
        '的过拟合迹象。Accuracy曲线在3~4个epoch后趋于平稳，后期的涨落主要由噪声样本的'
        '随机性造成。')

    heading(doc, '3.3 最终评估', 2)
    body(doc,
        '在1,000张验证集上的最终结果：')

    table(doc,
        ['类别', 'Precision', 'Recall', 'F1-Score', 'Support'],
        [
            ['圆形 (0)', '0.9979', '0.9897', '0.9938', '484'],
            ['方形 (1)', '0.9904', '0.9981', '0.9942', '516'],
            ['Accuracy', '', '', '0.9940', '1000'],
            ['Macro Avg', '0.9942', '0.9939', '0.9940', '1000'],
            ['Weighted Avg', '0.9940', '0.9940', '0.9940', '1000'],
        ]
    )

    body(doc,
        '验证集中圆形484张、方形516张，类别基本均衡。圆形召回率为0.9897（484张中5张被错判'
        '为方形），方形精确率为0.9904（516张中正确识别515张，另有1张圆形被误判为方形导致'
        '精确率下降）。总体准确率99.40%，共6张样本分类错误。相比噪声水平0.15时仅1个错误，'
        '噪声提升到0.25后错误数增加了5个，验证Loss从约0.003升至约0.032，说明高噪声显著增加'
        '了分类难度。')

    fig_placeholder(doc, '图4 验证集混淆矩阵', 'confusion_matrix.png')

    body(doc,
        '混淆矩阵（图4）直观展示了误判分布：6个错误中有5个是圆形被误判为方形，仅1个是方形'
        '被误判为圆形。圆形更容易被误判的原因是：高噪声条件下，部分圆形样本的轮廓被随机噪声'
        '打断，产生了类似角点的局部特征，CNN容易将其误读为方形的转折点信号。方形相对更难被'
        '误判（召回率0.9981），因为即使噪声严重，四条直线的组合特征仍比光滑曲线更稳定。这个'
        '不对称的误判模式可以作为后续改进模型的切入点。')

    heading(doc, '3.4 预测可视化', 2)

    fig_placeholder(doc, '图5 模型在验证集上的预测结果（8张随机样本）', 'predictions.png')

    body(doc,
        '图5随机选取了8张验证集样本进行可视化，每张图片下方标注真实类别和预测结果（绿色'
        '表示正确，红色表示错误）。可以看到大部分样本预测正确，模型对不同位置、不同大小和'
        '不同旋转角度的形状都具有较好的泛化能力。')

    heading(doc, '3.5 噪声水平对比实验', 2)
    body(doc,
        '表3-5对比了两种噪声水平下的模型表现：')

    table(doc,
        ['噪声水平 σ', '训练准确率(%)', '验证准确率(%)', '验证Loss', '错误样本数'],
        [
            ['0.05', '100.00', '100.00', '~0.0001', '0'],
            ['0.15', '100.00', '99.90', '~0.003', '1'],
            ['0.25', '99.88', '99.40', '~0.032', '6'],
        ]
    )

    body(doc,
        '噪声从0.05增加到0.25后，验证准确率从100%降至99.40%，验证Loss从约0.0001增大到'
        '约0.032，增大了两个数量级。错误样本数从0增加到6。噪声越大，圆形和方形的边缘越模糊，'
        'CNN提取的特征质量下降，分类确定性降低（体现为Loss上升）。这验证了卷积核对边缘特征'
        '的依赖性：当边缘信息被噪声掩盖时，分类性能下降。在实际应用中，可以通过以下方式提升'
        '抗噪能力：（1）增加训练数据量，覆盖更多的噪声模式；（2）使用更深的网络或更大的卷积'
        '核捕获更多上下文信息；（3）在训练时加入更强的数据增强，迫使模型学习更鲁棒的特征。')

    # ======== 四、讨论 ========
    heading(doc, '四、讨论', 1)

    heading(doc, '4.1 为什么CNN能有效区分圆形和方形', 2)
    body(doc,
        '从卷积核的功能角度分析：第一层卷积核通常学习边缘检测器（如Gabor滤波器），对图像'
        '中的边缘方向和强度敏感。圆形的边缘是光滑曲线，在所有方向上曲率一致；方形的边缘'
        '由四段直线和四个直角转折点组成。随着卷积层的加深，第二、三层卷积核组合底层边缘'
        '信息，形成对"曲率分布"和"角点"的检测能力。方形的四个角点会在特征图中产生显著的'
        '激活响应，这是CNN区分两种形状的关键依据。')

    heading(doc, '4.2 本实验的局限性', 2)
    body(doc,
        '一是任务偏简单。圆形和方形在生成时属于理想几何图形，形态差异明显，即使简单的机器'
        '学习方法（如HOG+SVM）也可能达到类似准确率。但本实验选用简单任务是有意为之——目的'
        '是观察CNN在"确定可解"问题上各组件（BN、Dropout、卷积深度）的作用，而非追求任务'
        '难度。')

    body(doc,
        '二是数据集缺乏多样性。所有图片由程序生成，缺少真实世界图像中的光照变化、背景干扰、'
        '部分遮挡等因素。因此模型虽然在本数据集上表现优异，但无法直接推广到自然图像中的'
        '形状识别。这可以通过更换为真实图像数据集或引入更复杂的合成数据来改进。')

    body(doc,
        '三是未与其他方法做系统对比。实验中只对比了不同噪声水平的结果，未与传统方法（如SVM'
        '基于HOG特征、模板匹配等）做横向比较，也未尝试不同CNN架构（如增加残差连接、'
        '使用不同激活函数等）。完整的对比实验可以作为后续工作的方向。')

    heading(doc, '4.3 可改进的方向', 2)
    body(doc,
        '（1）引入更多几何形状类别（三角形、五边形等），将二分类扩展为多分类任务；'
        '（2）使用Grad-CAM等可视化工具展示CNN的关注区域，直观验证模型是否确实关注形状轮廓；'
        '（3）在训练时加入更强的数据扰动（弹性形变、部分遮挡、随机背景纹理），测试模型的'
        '鲁棒性上限；'
        '（4）将模型部署到简单的GUI应用中，支持用户手绘图形并实时预测。')

    # ======== 五、总结 ========
    heading(doc, '五、总结', 1)
    body(doc,
        '本实验使用PyTorch搭建了一个三层卷积的CNN，对程序生成的圆形与方形图像进行二分类。'
        '在噪声水平σ=0.25的条件下，模型在5,000张训练样本上训练15个epoch后，在1,000张验证'
        '样本上取得了99.40%的分类准确率（6个错误）。实验结果表明：CNN能够有效提取几何形状的'
        '边缘和轮廓'
        '特征，Batch Normalization和Dropout对稳定训练和防止过拟合起到了积极作用。噪声水平的'
        '提高会导致分类性能下降，原因是边缘信息被噪声破坏后，卷积核提取的特征质量降低。'
        '本实验的不足之处在于任务较为简单、缺乏与基线方法的系统对比，以及数据集与真实场景'
        '存在差异。在后续工作中可以扩展类别数量、引入可解释性分析工具，并结合实际图像数据'
        '进一步验证CNN的泛化能力。')

    # ======== 参考文献 ========
    heading(doc, '参考文献', 1)
    refs = [
        '[1] LeCun Y, Bottou L, Bengio Y, et al. Gradient-based learning applied to document recognition[J]. '
        'Proceedings of the IEEE, 1998, 86(11): 2278-2324.',
        '[2] Ioffe S, Szegedy C. Batch normalization: Accelerating deep network training by reducing internal '
        'covariate shift[C]. ICML, 2015: 448-456.',
        '[3] Srivastava N, Hinton G, Krizhevsky A, et al. Dropout: A simple way to prevent neural networks '
        'from overfitting[J]. Journal of Machine Learning Research, 2014, 15(1): 1929-1958.',
        '[4] Kingma D P, Ba J. Adam: A method for stochastic optimization[C]. ICLR, 2015.',
        '[5] PyTorch Documentation. https://pytorch.org/docs/stable/index.html',
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.line_spacing = 1.5
        r = p.add_run(ref)
        set_run_font(r, '宋体', 'Times New Roman', Pt(11))

    # ===== 保存 =====
    out = 'CNN课程报告_圆形vs方形分类.docx'
    doc.save(out)
    print(f'[完成] {out}')
    print('图片待插入位置已在文档中用 [插入图片: xxx.png] 标注。')


if __name__ == '__main__':
    generate()
